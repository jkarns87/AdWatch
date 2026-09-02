"""DOCX renderer (python-docx) — for readers who want to edit the brief before forwarding it."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ASSETS = Path(__file__).parent / "assets"
NAVY = RGBColor(0x0B, 0x10, 0x20)
MUTED = RGBColor(0x6B, 0x73, 0x90)


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def _table(doc, header: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, wdt in enumerate(widths):
                row.cells[i].width = Inches(wdt)
    return t


def render_docx(data: dict[str, Any]) -> bytes:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    logo = ASSETS / "logo-light.png"
    if logo.exists():
        doc.add_picture(str(logo), width=Inches(1.8))
    t = doc.add_paragraph()
    r = t.add_run(data["title"])
    r.font.size, r.font.bold, r.font.color.rgb = Pt(22), True, NAVY
    m = doc.add_paragraph()
    mr = m.add_run(
        f"{data['watchlist']['name']} · {data['watchlist']['vertical']} · {data['watchlist']['geo']}   |   "
        f"{data['period']['from']} → {data['period']['to']}   |   prepared for {'the CFO' if data['audience'] == 'cfo' else 'Marketing'}"
    )
    mr.font.size, mr.font.color.rgb = Pt(9), MUTED

    k = data["kpis"]
    kp = doc.add_paragraph()
    kp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kr = kp.add_run(
        f"{k['changes']} changes ({k['high']} high · {k['medium']} medium · {k['low']} low)   ·   {k['insights']} AI insights   ·   "
        f"{k['actions']} actions   ·   {k['competitors']} competitors   ·   {k['keywords']} keywords   ·   {k['runs_in_period']} runs"
    )
    kr.font.size, kr.font.bold = Pt(9.5), True

    ex = data["executive_summary"]
    _h(doc, "Executive summary")
    hp = doc.add_paragraph()
    hr = hp.add_run(ex.get("headline", ""))
    hr.font.bold, hr.font.size = True, Pt(12)
    for p in ex.get("paragraphs", []):
        doc.add_paragraph(p)
    if ex.get("decisions"):
        doc.add_paragraph("Decisions for this reader").runs[0].font.bold = True
        for d in ex["decisions"]:
            doc.add_paragraph(d, style="List Bullet")
    if ex.get("watch_next"):
        wp = doc.add_paragraph()
        wp.add_run("Watch next: ").font.bold = True
        wp.add_run(ex["watch_next"])

    if data["actions"]:
        _h(doc, "Recommended actions")
        _table(doc, ["Urgency", "Effort", "Action", "Rationale"], [[a.get("urgency", "").replace("_", " "), a.get("effort", ""), a.get("action", ""), a.get("rationale", "")] for a in data["actions"]], [0.9, 0.7, 2.6, 2.8])

    _h(doc, f"What changed ({k['changes']} events)")
    if data["changes_by_kind"]:
        s = doc.add_paragraph(" · ".join(f"{c['count']} {c['label'].lower()}" for c in data["changes_by_kind"]))
        s.runs[0].font.size, s.runs[0].font.color.rgb = Pt(9), MUTED
    _table(doc, ["Severity", "Type", "Subject", "Detail"], [[c["severity"], c["label"], c["subject"], c["description"]] for c in data["changes"]] or [["—", "No changes detected in the period", "", ""]], [0.7, 1.4, 1.3, 3.6])

    _h(doc, "Competitor activity")
    _table(doc, ["Competitor", "Domain", "Active creatives", "Launched", "Dropped", "Formats"], [[c["name"], c["domain"], c["active_creatives"], c["launched"], c["dropped"], ", ".join(f"{a} {b}" for a, b in sorted(c["formats"].items()))] for c in data["competitors"]])

    _h(doc, "Keywords: paid block, share of voice, demand")
    for kw in data["keywords"]:
        p = doc.add_paragraph()
        p.add_run(kw["term"]).font.bold = True
        block = ", ".join(f"{'↓' if a['block'] == 'bottom' else ''}#{a['position']} {a['domain']}{' *' if a['tracked'] else ''}" for a in kw["paid_block"]) or "no paid results captured"
        sov = ", ".join(f"{s['domain']} ×{s['appearances']} (avg #{s['avg_position']})" for s in kw["share_of_voice"][:4]) or "—"
        demand = (f"{kw['demand_latest']}" + (f" vs 4-wk avg {kw['demand_trailing']}" if kw["demand_trailing"] is not None else "")) if kw["demand_latest"] is not None else "—"
        for label, val in (("Paid block now", block), ("Share of voice (all runs)", sov), ("Demand", demand), ("Rising", ", ".join(kw["rising_queries"]) or "—")):
            q = doc.add_paragraph()
            q.paragraph_format.space_after = Pt(0)
            q.add_run(f"{label}: ").font.bold = True
            q.add_run(val)
    f = doc.add_paragraph("* tracked competitor. Source: SerpApi (Google Ads Transparency Center, Google Search, Google Trends). Analysis: AdWatch diff engine + AI analyst.")
    f.runs[0].font.size, f.runs[0].font.color.rgb = Pt(8), MUTED

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
